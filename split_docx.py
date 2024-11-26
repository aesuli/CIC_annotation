import os
import re

from docx import Document

if __name__ == '__main__':
    data_dir = './data'
    out_doc_dir = f'{data_dir}/docx'
    out_txt_dir = f'{data_dir}/txt'
    os.makedirs(out_doc_dir, exist_ok=True)
    os.makedirs(out_txt_dir,exist_ok=True)

    with open(f'{data_dir}/lemma_glossato.txt',mode='wt',encoding='utf-8') as lemma_file:
        doc = Document(f'{data_dir}/Decretals Gloss, Books 1-5 Complete, rev. 9.23.docx')
        newdoc = Document()
        text = ''
        name = None
        re_chapt = re.compile('(REX PACIFICUS)|X ([0-9]+\.[0-9]+ .*)')
        for paragraph in doc.paragraphs:
            match = re_chapt.match(paragraph.text)
            if match:
                if name:
                    newdoc.save(f'{out_doc_dir}/{name}.docx')
                    with open(f'{out_txt_dir}/{name}.txt', mode='wt',encoding='utf-8') as output_file:
                        output_file.write(text)
                newdoc = Document()
                text = ''
                name = match.group(1)
                if name is None:
                    name = match.group(2)
                print(name)
            newdoc.add_paragraph(paragraph.text, style=paragraph.style)
            prefix = ''
            if paragraph.style.name == 'Heading 1':
                prefix = ''
            elif paragraph.style.name == 'Heading 2':
                prefix = ''
            elif paragraph.style.name == 'Heading 3':
                prefix = ''
            elif paragraph.style.name == 'Heading 4':
                prefix = ' '
                print(f'{paragraph.text}',file=lemma_file)
            elif paragraph.style.name != 'Normal':
                print(paragraph.style.name)
            text += f'{prefix}{paragraph.text}\n'
        if name:
            newdoc.save(f'{out_doc_dir}/{name}.docx')
            with open(f'{out_txt_dir}/{name}.txt', mode='wt', encoding='utf-8') as output_file:
                output_file.write(text)
